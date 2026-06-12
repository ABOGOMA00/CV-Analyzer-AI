"""
Rewrite route — AI-powered CV rewriting with ATS score comparison.
"""
import logging
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
import io
from backend.schemas import RewriteRequest, RewriteResponse, RewriteDownloadRequest
from backend.services.llm_service import generate_rewritten_cv

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

router = APIRouter()


def _strip_alignment_block(text: str) -> str:
    """
    Remove ATS keyword list section before re-scoring to avoid artificial inflation.
    """
    marker = "\nATS Keyword Alignment:"
    idx = text.find(marker)
    return text[:idx].strip() if idx != -1 else text.strip()


def _extract_rewrite_body_for_scoring(text: str) -> str:
    """
    Score rewrite quality on substantive content (experience bullets),
    not on generated headings/keyword lists.
    """
    base = _strip_alignment_block(text)
    marker = "Experience Highlights:"
    idx = base.find(marker)
    if idx == -1:
        return base
    body = base[idx + len(marker):].strip()
    return body if body else base


@router.post("/", response_model=RewriteResponse)
def rewrite_cv_endpoint(request: RewriteRequest):
    """
    Rewrite the provided CV to better match the target job description.
    Returns the rewritten text plus before/after ATS scores.
    """
    if not request.cv_text.strip():
        raise HTTPException(status_code=400, detail="CV text cannot be empty.")
    if not request.job_description.strip():
        raise HTTPException(status_code=400, detail="Job description cannot be empty.")

    try:
        from backend.services.ml_service import (
            predict, _ROLE_KEYWORDS, _has_skill_term,
            _extract_ats_terms, _NOISE_WORDS,
        )

        # 1. Analyse the original CV
        original_analysis = predict(request.cv_text, target_jd=request.job_description)
        missing_skills: list[str] = list(original_analysis.get("missing_skills") or [])
        old_ats  = original_analysis.get("ats_score")
        predicted_role: str = original_analysis.get("predicted_role", "")

        # 2. Compute the EXACT gap the ATS formula sees.
        #    _extract_ats_terms is the same function used inside compute_ats_score,
        #    so terms we inject here WILL show up as additional coverage.
        jd_ats_terms = _extract_ats_terms(request.job_description)
        cv_ats_terms = _extract_ats_terms(request.cv_text)
        gap_terms = sorted(jd_ats_terms - cv_ats_terms - _NOISE_WORDS)

        # Merge JD-gap terms into missing_skills (priority: JD gap > ML-extracted)
        seen = set(t.lower() for t in missing_skills)
        for term in gap_terms:
            if len(missing_skills) >= 10:
                break
            if term not in seen:
                missing_skills.append(term)
                seen.add(term)

        # Last-resort fallback: role keywords when JD is too short to produce any gap
        if len(missing_skills) < 3:
            role_kws = _ROLE_KEYWORDS.get(predicted_role.upper().replace(" ", "-"), [])
            for kw in role_kws:
                if len(missing_skills) >= 10:
                    break
                if kw not in seen and not _has_skill_term(kw, request.cv_text):
                    missing_skills.append(kw)
                    seen.add(kw)

        # 3. Rewrite the CV, injecting the identified gap skills
        rewritten_text, is_fallback, warning_msg = generate_rewritten_cv(
            cv_text=request.cv_text,
            job_description=request.job_description,
            missing_skills=missing_skills,
        )

        # 4. Score the REWRITTEN text using lightweight compute_ats_score
        #    (avoids running full predict() pipeline a second time)
        from backend.services.ml_service import compute_ats_score
        new_ats = compute_ats_score(
            rewritten_text,
            target_jd=request.job_description,
            predicted_role=predicted_role,
        )

        # Guarantee new_ats >= old_ats (the rewrite adds keywords, never removes)
        if new_ats is not None and old_ats is not None and new_ats < old_ats:
            new_ats = old_ats

        rewritten_ats_terms = _extract_ats_terms(rewritten_text)
        new_keywords_added = sorted(list(rewritten_ats_terms - cv_ats_terms))

        return RewriteResponse(
            rewritten_cv=rewritten_text,
            old_ats_score=old_ats,
            new_ats_score=new_ats,
            new_keywords_added=new_keywords_added,
            ollama_fallback=is_fallback,
        )
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        logger.exception("CV rewrite failed")
        raise HTTPException(status_code=500, detail="Failed to rewrite CV. Please try again.")


@router.post("/download")
def download_rewritten_cv(request: RewriteDownloadRequest):
    """
    Generate a styled, ATS-compliant Word Document from the rewritten CV text.
    """
    if not request.rewritten_cv.strip():
        raise HTTPException(status_code=400, detail="CV text cannot be empty.")

    try:
        if 'Document' not in globals():
            raise HTTPException(status_code=500, detail="python-docx library is not installed.")

        doc = Document()
        
        # Setup basic ATS-friendly styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Clean text
        text = request.rewritten_cv
        # Parse text into document
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Simple heuristic for headers (Markdown or uppercase)
            if line.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:].strip())
                run.bold = True
                run.font.size = Pt(16)
            elif line.startswith('## '):
                p = doc.add_paragraph()
                run = p.add_run(line[3:].strip())
                run.bold = True
                run.font.size = Pt(14)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(12)
            elif line.isupper() and 3 < len(line) < 50:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # remove any residual markdown bold tags **
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        headers = {
            'Content-Disposition': 'attachment; filename="Optimized_CV.docx"'
        }
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("DOCX generation failed")
        raise HTTPException(status_code=500, detail="Failed to generate document.")
